"""
Document parser service for extracting text from PDF, DOCX, and EML files.
"""
import io
import re
import email
from email import policy
from email.parser import BytesParser
from typing import Optional, Tuple
from datetime import datetime

from PyPDF2 import PdfReader
from docx import Document


class DocumentParseError(Exception):
    """Raised when document parsing fails."""
    pass


class DocumentParserService:
    """Service for parsing various document formats."""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.eml'}

    def get_file_type(self, filename: str) -> str:
        """Determine content type based on file extension."""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        type_map = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',
            'eml': 'email',
        }
        return type_map.get(ext, 'document')

    def parse_file(self, file_content: bytes, filename: str) -> Tuple[str, str, dict]:
        """
        Parse a file and extract text content.

        Args:
            file_content: Raw bytes of the file
            filename: Original filename (used to determine type)

        Returns:
            Tuple of (title, text_content, metadata)
        """
        ext = filename.lower().split('.')[-1] if '.' in filename else ''

        if ext == 'pdf':
            return self._parse_pdf(file_content, filename)
        elif ext in ('docx', 'doc'):
            return self._parse_docx(file_content, filename)
        elif ext == 'eml':
            return self._parse_eml(file_content, filename)
        else:
            raise DocumentParseError(f"Unsupported file type: {ext}")

    def _parse_pdf(self, content: bytes, filename: str) -> Tuple[str, str, dict]:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(io.BytesIO(content))

            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            full_text = '\n\n'.join(text_parts)

            if not full_text.strip():
                raise DocumentParseError("PDF contains no extractable text (may be scanned/image-based)")

            # Try to get title from metadata or use filename
            title = filename.rsplit('.', 1)[0]
            if reader.metadata:
                pdf_title = reader.metadata.get('/Title')
                if pdf_title and pdf_title.strip():
                    title = pdf_title.strip()

            # Get first line as potential title if it's short
            first_line = full_text.strip().split('\n')[0].strip()
            if len(first_line) < 100 and first_line:
                title = first_line

            metadata = {
                'source': 'file_upload',
                'file_type': 'pdf',
                'original_filename': filename,
                'page_count': len(reader.pages),
            }

            # Add PDF metadata if available
            if reader.metadata:
                if reader.metadata.get('/Author'):
                    metadata['author'] = reader.metadata.get('/Author')
                if reader.metadata.get('/CreationDate'):
                    metadata['creation_date'] = str(reader.metadata.get('/CreationDate'))

            return title, self._clean_text(full_text), metadata

        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse PDF: {str(e)}")

    def _parse_docx(self, content: bytes, filename: str) -> Tuple[str, str, dict]:
        """Extract text from Word document."""
        try:
            doc = Document(io.BytesIO(content))

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            full_text = '\n\n'.join(text_parts)

            if not full_text.strip():
                raise DocumentParseError("Document contains no text")

            # Try to get title
            title = filename.rsplit('.', 1)[0]

            # Use first heading or first paragraph as title
            for para in doc.paragraphs:
                if para.text.strip():
                    if len(para.text.strip()) < 100:
                        title = para.text.strip()
                    break

            metadata = {
                'source': 'file_upload',
                'file_type': 'docx',
                'original_filename': filename,
                'paragraph_count': len(doc.paragraphs),
            }

            # Try to get core properties
            try:
                if doc.core_properties:
                    if doc.core_properties.author:
                        metadata['author'] = doc.core_properties.author
                    if doc.core_properties.created:
                        metadata['creation_date'] = doc.core_properties.created.isoformat()
                    if doc.core_properties.title:
                        title = doc.core_properties.title
            except:
                pass

            return title, self._clean_text(full_text), metadata

        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse Word document: {str(e)}")

    def _parse_eml(self, content: bytes, filename: str) -> Tuple[str, str, dict]:
        """Extract text from email (.eml) file."""
        try:
            msg = BytesParser(policy=policy.default).parsebytes(content)

            # Get subject as title
            subject = msg.get('Subject', '')
            title = subject if subject else filename.rsplit('.', 1)[0]

            # Get sender and date
            sender = msg.get('From', '')
            date_str = msg.get('Date', '')

            # Extract body
            body_parts = []

            # Add header info
            header_info = []
            if sender:
                header_info.append(f"De: {sender}")
            if msg.get('To'):
                header_info.append(f"Para: {msg.get('To')}")
            if date_str:
                header_info.append(f"Fecha: {date_str}")
            if subject:
                header_info.append(f"Asunto: {subject}")

            if header_info:
                body_parts.append('\n'.join(header_info))
                body_parts.append('---')

            # Get email body
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == 'text/plain':
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or 'utf-8'
                            try:
                                body_parts.append(payload.decode(charset))
                            except:
                                body_parts.append(payload.decode('utf-8', errors='ignore'))
                    elif content_type == 'text/html' and not any('text/plain' in str(p.get_content_type()) for p in msg.walk() if p != part):
                        # Only use HTML if no plain text available
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or 'utf-8'
                            try:
                                html_content = payload.decode(charset)
                                # Simple HTML to text conversion
                                body_parts.append(self._html_to_text(html_content))
                            except:
                                pass
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    charset = msg.get_content_charset() or 'utf-8'
                    try:
                        body_parts.append(payload.decode(charset))
                    except:
                        body_parts.append(payload.decode('utf-8', errors='ignore'))

            full_text = '\n\n'.join(body_parts)

            if not full_text.strip():
                raise DocumentParseError("Email contains no text content")

            metadata = {
                'source': 'file_upload',
                'file_type': 'email',
                'original_filename': filename,
                'email_subject': subject,
                'email_from': sender,
                'email_to': msg.get('To', ''),
                'email_date': date_str,
            }

            return title, self._clean_text(full_text), metadata

        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse email: {str(e)}")

    def _html_to_text(self, html: str) -> str:
        """Simple HTML to plain text conversion."""
        # Remove scripts and styles
        text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)

        # Convert line breaks
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</p>', '\n\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</li>', '\n', text, flags=re.IGNORECASE)

        # Remove all remaining HTML tags
        text = re.sub(r'<[^>]+>', '', text)

        # Decode common HTML entities
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&quot;', '"')
        text = text.replace('&#39;', "'")

        return text

    def _clean_text(self, text: str) -> str:
        """Clean up extracted text."""
        # Normalize whitespace
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)

        # Remove excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text.strip()


# Singleton instance
document_parser = DocumentParserService()
